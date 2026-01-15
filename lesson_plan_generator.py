import streamlit as st
import json
import io
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import zipfile
import tempfile
from pathlib import Path

# Set page config
st.set_page_config(
    page_title="Plan Filler: Weekly Lesson Plan Generator",
    page_icon="📚",
    layout="wide"
)

# Custom CSS
st.markdown("""
<style>
    .main-title {
        text-align: center;
        color: #1E3A8A;
        font-size: 2.5em;
        margin-bottom: 10px;
    }
    .subtitle {
        text-align: center;
        color: #4B5563;
        font-size: 1.2em;
        margin-bottom: 40px;
    }
    .upload-box {
        border: 2px dashed #4F46E5;
        border-radius: 10px;
        padding: 30px;
        text-align: center;
        background-color: #F8FAFC;
        margin: 20px 0;
    }
    .success-box {
        background-color: #D1FAE5;
        padding: 20px;
        border-radius: 10px;
        border-left: 5px solid #10B981;
        margin: 20px 0;
    }
    .info-box {
        background-color: #E0F2FE;
        padding: 20px;
        border-radius: 10px;
        border-left: 5px solid #0EA5E9;
        margin: 20px 0;
    }
    .stDownloadButton > button {
        width: 100% !important;
        margin: 5px 0 !important;
        background-color: #4F46E5 !important;
        color: white !important;
    }
    .template-info {
        font-size: 0.9em;
        color: #6B7280;
        font-style: italic;
    }
</style>
""", unsafe_allow_html=True)

# App title
st.markdown('<h1 class="main-title">📚 Plan Filler: Weekly Lesson Plan Generator</h1>', unsafe_allow_html=True)
st.markdown('<p class="subtitle">Upload your Word template or use the default template, then paste the JSON of your weekly lesson plan. Click Generate Lesson Plan to get a ready-to-use Word document.</p>', unsafe_allow_html=True)

# JSON skeleton for reference - updated to match Word template placeholders
JSON_SKELETON = '''{
  "Teacher": "Teacher Name Here",
  "WeekNumber": "Week Number Here",
  "YearClass": "Year/Class Here (e.g., 7)",
  "Subject": "Subject Here (e.g., English)",
  "UnitTopic": "Unit/Topic Here (e.g., Adventure Stories)",

  "Schedule": {
    "Day1": "Sunday - Class 1 (Single)",
    "Day2": "Monday - Class 2 (Single)",
    "Day3": "Tuesday/Thursday - Class 3 (Double) - 7B(Tuesday) / 7A(Thursday)",
    "Day4": "Wednesday - Class 4 (Double)",
    "Day5": "Tuesday/Thursday - Class 5 (Single) - 7A(Tuesday) / 7B(Thursday)"
  },

  "LearningObjective": {
    "Class1_LearningObjective": "",
    "Class2_LearningObjective": "",
    "Class3_LearningObjective": "",
    "Class4_LearningObjective": "",
    "Class5_LearningObjective": ""
  },

  "SuccessCriteria": {
    "Class1_SuccessCriteria": "",
    "Class2_SuccessCriteria": "",
    "Class3_SuccessCriteria": "",
    "Class4_SuccessCriteria": "",
    "Class5_SuccessCriteria": ""
  },

  "KeyVocabulary": {
    "Class1_Vocabulary": "",
    "Class2_Vocabulary": "",
    "Class3_Vocabulary": "",
    "Class4_Vocabulary": "",
    "Class5_Vocabulary": ""
  },

  "KeyQuestions": {
    "Class1_KeyQuestions": "",
    "Class2_KeyQuestions": "",
    "Class3_KeyQuestions": "",
    "Class4_KeyQuestions": "",
    "Class5_KeyQuestions": ""
  },

  "StarterActivity": {
    "Class1_StarterActivity": "",
    "Class2_StarterActivity": "",
    "Class3_StarterActivity": "",
    "Class4_StarterActivity": "",
    "Class5_StarterActivity": ""
  },

  "TeacherInputMainTeaching": {
    "Class1_TeacherInputMainTeaching": "",
    "Class2_TeacherInputMainTeaching": "",
    "Class3_TeacherInputMainTeaching": "",
    "Class4_TeacherInputMainTeaching": "",
    "Class5_TeacherInputMainTeaching": ""
  },

  "DifferentiatedActivities": {
    "Class1_DifferentiatedActivities": "",
    "Class2_DifferentiatedActivities": "",
    "Class3_DifferentiatedActivities": "",
    "Class4_DifferentiatedActivities": "",
    "Class5_DifferentiatedActivities": ""
  },

  "Plenary": {
    "Class1_Plenary": "",
    "Class2_Plenary": "",
    "Class3_Plenary": "",
    "Class4_Plenary": "",
    "Class5_Plenary": ""
  },

  "Reflection": {
    "Class1_Reflection": "",
    "Class2_Reflection": "",
    "Class3_Reflection": "",
    "Class4_Reflection": "",
    "Class5_Reflection": ""
  },

  "Homework": {
    "Class1_Homework": "",
    "Class2_Homework": "",
    "Class3_Homework": "",
    "Class4_Homework": "",
    "Class5_Homework": ""
  }
}'''


def create_default_template():
    """Create a default WPLT template"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Weekly Lesson Plan Template', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Teacher information
    doc.add_paragraph(f"Teacher: {{Teacher}}")
    doc.add_paragraph(f"Week number: {{WeekNumber}}")
    doc.add_paragraph(f"Year/Class: {{YearClass}}")
    doc.add_paragraph(f"Subject: {{Subject}}")
    doc.add_paragraph(f"Unit / Topic: {{UnitTopic}}")
    
    # Schedule
    doc.add_paragraph()
    doc.add_heading('Schedule', level=1)
    
    schedule_data = [
        ("Class 1 (Single)", "{{Schedule_Day1}}"),
        ("Class 2 (Single)", "{{Schedule_Day2}}"),
        ("Class 3 (Double)", "{{Schedule_Day3}}"),
        ("Class 4 (Double)", "{{Schedule_Day4}}"),
        ("Class 5 (Single)", "{{Schedule_Day5}}")
    ]
    
    for class_info, day_info in schedule_data:
        doc.add_paragraph(f"{class_info}: {day_info}")
    
    # Lesson Plan Sections
    sections = [
        ("Learning Objective", "LearningObjective"),
        ("Success Criteria", "SuccessCriteria"),
        ("Key Vocabulary", "KeyVocabulary"),
        ("Key Questions", "KeyQuestions"),
        ("Starter Activity (hook)", "StarterActivity"),
        ("Teacher input / Main Teaching", "TeacherInputMainTeaching"),
        ("Differentiated Activities", "DifferentiatedActivities"),
        ("Plenary", "Plenary"),
        ("Reflection", "Reflection"),
        ("Homework", "Homework")
    ]
    
    for section_name, section_key in sections:
        doc.add_paragraph()
        doc.add_heading(section_name, level=1)
        
        # Create a table for each class
        table = doc.add_table(rows=2, cols=6)
        table.style = 'Table Grid'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = section_name
        for i in range(1, 6):
            header_cells[i].text = f"Class {i}"
        
        # Data row with placeholders
        data_cells = table.rows[1].cells
        data_cells[0].text = section_name
        for i in range(1, 6):
            data_cells[i].text = f"{{{{Class{i}_{section_key}}}}}"
    
    return doc

def fill_template_with_json(template_doc, json_data):
    """Fill the template with JSON data"""
    
    # Convert JSON string to dict if needed
    if isinstance(json_data, str):
        try:
            data = json.loads(json_data)
        except json.JSONDecodeError as e:
            return None, f"Invalid JSON: {str(e)}"
    else:
        data = json_data
    
    # Create a mapping of placeholders to values
    placeholder_map = {}
    
    # Basic info
    basic_fields = ['Teacher', 'WeekNumber', 'YearClass', 'Subject', 'UnitTopic']
    for field in basic_fields:
        if field in data:
            placeholder_map[f'{{{field}}}'] = str(data[field])
    
    # Schedule
    if 'Schedule' in data:
        schedule = data['Schedule']
        for i in range(1, 6):
            day_key = f'Day{i}'
            if day_key in schedule:
                placeholder_map[f'{{Schedule_{day_key}}}'] = str(schedule[day_key])
    
    # Lesson sections
    sections = [
        'LearningObjective',
        'SuccessCriteria', 
        'KeyVocabulary',
        'KeyQuestions',
        'StarterActivity',
        'TeacherInputMainTeaching',
        'DifferentiatedActivities',
        'Plenary',
        'Reflection',
        'Homework'
    ]
    
    for section in sections:
        if section in data:
            section_data = data[section]
            for i in range(1, 6):
                class_key = f'Class{i}_{section}'
                if class_key in section_data:
                    placeholder_map[f'{{{{Class{i}_{section}}}}}'] = str(section_data[class_key])
    
    # Replace placeholders in the document
    for paragraph in template_doc.paragraphs:
        for placeholder, value in placeholder_map.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)
    
    # Replace in tables
    for table in template_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, value in placeholder_map.items():
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, value)
    
    # Replace in headers
    for section in template_doc.sections:
        for paragraph in section.header.paragraphs:
            for placeholder, value in placeholder_map.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, value)
    
    return template_doc, "Success"

def save_document_to_bytes(doc):
    """Save document to bytes"""
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Main app layout
col1, col2 = st.columns([2, 1])

with col1:
    # Upload template section
    st.markdown("### 📁 Upload your Lesson Plan Template (.docx)")
    
    uploaded_file = st.file_uploader(
        "Drag and drop file here",
        type=['docx'],
        help="Upload your school's Word template or use the default template below",
        label_visibility="collapsed"
    )
    
    st.caption("Limit 200MB per file • DOCX")
    
    # Display template status
    if uploaded_file:
        st.success("✅ Template uploaded successfully!")
        template_bytes = uploaded_file.read()
        template_doc = Document(io.BytesIO(template_bytes))
    else:
        st.info("No template uploaded or found in repo. Using default WPLT template.")
        template_doc = create_default_template()
        template_bytes = save_document_to_bytes(template_doc).getvalue()
    
    # JSON input section
    st.markdown("---")
    st.markdown("### 📝 Paste your Lesson Plan JSON")
    
    # Sample JSON button
    if st.button("📋 Load Sample JSON", key="load_sample"):
        st.session_state.json_input = JSON_SKELETON
    
    json_input = st.text_area(
        "Paste your JSON here:",
        height=400,
        value=st.session_state.get('json_input', JSON_SKELETON),
        help="Use the JSON skeleton format above"
    )
    
    # Generate button
    if st.button("🚀 Generate Lesson Plan", type="primary", use_container_width=True):
        if json_input:
            try:
                # Validate JSON
                data = json.loads(json_input)
                
                # Check required fields
                required_sections = ['LearningObjective', 'SuccessCriteria']
                missing = [section for section in required_sections if section not in data]
                
                if missing:
                    st.error(f"Missing required sections: {', '.join(missing)}")
                else:
                    # Create progress indicator
                    progress_bar = st.progress(0)
                    
                    with st.spinner("Processing your lesson plan..."):
                        # Reload template (in case it was modified)
                        if uploaded_file:
                            template_doc = Document(io.BytesIO(template_bytes))
                        else:
                            template_doc = create_default_template()
                        
                        progress_bar.progress(30)
                        
                        # Fill template with JSON data
                        filled_doc, message = fill_template_with_json(template_doc, data)
                        
                        progress_bar.progress(70)
                        
                        if filled_doc:
                            # Save to bytes
                            output_buffer = save_document_to_bytes(filled_doc)
                            progress_bar.progress(100)
                            
                            # Success message
                            st.markdown('<div class="success-box">✅ Lesson plan generated successfully! Download your file below.</div>', unsafe_allow_html=True)
                            
                            # Get filename from JSON or use default
                            teacher_name = data.get('Teacher', 'Teacher').replace(' ', '_')
                            week_num = data.get('WeekNumber', 'Week')
                            subject = data.get('Subject', 'Subject')
                            filename = f"{teacher_name}_{week_num}_{subject}_Lesson_Plan.docx"
                            
                            # Download button
                            st.download_button(
                                label="📥 Download Lesson Plan (.docx)",
                                data=output_buffer,
                                file_name=filename,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                            
                            # Preview section
                            with st.expander("🔍 Preview JSON Data"):
                                st.json(data)
                                
                                # Show stats
                                col_stats1, col_stats2, col_stats3 = st.columns(3)
                                with col_stats1:
                                    st.metric("Teacher", data.get('Teacher', 'Not specified'))
                                with col_stats2:
                                    st.metric("Week", data.get('WeekNumber', 'Not specified'))
                                with col_stats3:
                                    st.metric("Subject", data.get('Subject', 'Not specified'))
                        else:
                            st.error(f"❌ Error: {message}")
                    
            except json.JSONDecodeError as e:
                st.error(f"❌ Invalid JSON format: {str(e)}")
            except Exception as e:
                st.error(f"❌ An error occurred: {str(e)}")
        else:
            st.warning("⚠️ Please paste your JSON data first.")

with col2:
    # Information panel
    st.markdown("### 📋 JSON Format Guide")
    
    with st.expander("📚 Click to see required format", expanded=True):
        st.markdown("""
        **Basic Information:**
        ```json
        "Teacher": "Your Name",
        "WeekNumber": "Week 5",
        "YearClass": "7",
        "Subject": "English",
        "UnitTopic": "Adventure Stories"
        ```
        
        **Schedule:**
        ```json
        "Schedule": {
          "Day1": "Sunday - Class 1 (Single)",
          "Day2": "Monday - Class 2 (Single)",
          "Day3": "Tuesday/Thursday - Class 3 (Double)",
          "Day4": "Wednesday - Class 4 (Double)",
          "Day5": "Tuesday/Thursday - Class 5 (Single)"
        }
        ```
        
        **For each lesson section (example):**
        ```json
        "LearningObjective": {
          "Class1_LearningObjective": "Students will learn...",
          "Class2_LearningObjective": "Students will understand...",
          "Class3_LearningObjective": "Students will be able to...",
          "Class4_LearningObjective": "Students will analyze...",
          "Class5_LearningObjective": "Students will create..."
        }
        ```
        
        **Available sections:**
        1. LearningObjective
        2. SuccessCriteria
        3. KeyVocabulary
        4. KeyQuestions
        5. StarterActivity
        6. TeacherInputMainTeaching
        7. DifferentiatedActivities
        8. Plenary
        9. Reflection
        10. Homework
        """)
    
    st.markdown("---")
    st.markdown("### 🎯 Quick Tips")
    
    st.markdown("""
    1. **Copy the skeleton** from the main area
    2. **Fill in your data** for all 5 classes
    3. **Use AI to convert** your text plans to JSON
    4. **Paste and generate** - that's it!
    
    **Template Notes:**
    - Uses {{Class1_LearningObjective}} format
    - Placeholders are automatically replaced
    - All 5 days are processed at once
    - Download ready-to-use Word file
    """)
    
    st.markdown("---")
    
    # Quick action buttons
    st.markdown("### ⚡ Quick Actions")
    
    col_btn1, col_btn2 = st.columns(2)
    
    with col_btn1:
        if st.button("📄 View Full Skeleton", use_container_width=True):
            st.code(JSON_SKELETON, language='json')
    
    with col_btn2:
        if st.button("🔄 Clear All", use_container_width=True):
            st.session_state.json_input = ""
            st.rerun()
    
    # Template info
    st.markdown("---")
    st.markdown('<p class="template-info">💡 <strong>Using default template</strong> - Upload your school template for custom formatting</p>', unsafe_allow_html=True)

# Footer
st.markdown("---")
st.caption("Plan Filler v1.0 • Automate your weekly lesson planning • No login required • Data stays in your browser")
